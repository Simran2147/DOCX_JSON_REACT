import json
import docx
from collections import defaultdict

def extract_data_from_docx(file_path):
    doc = docx.Document(file_path)
    extracted_data = {
        "headers": [],
        "content": defaultdict(str),
        "tables": []
    }
    
    current_section = None
    last_table_name = None
    paragraphs = list(doc.paragraphs)
    tables = doc.tables
    table_index = 0
    
    for i, para in enumerate(paragraphs):
        if para.style.name.startswith("Heading"):
            current_section = para.text.strip()
            if current_section!="":
                extracted_data["headers"].append(current_section)
            last_table_name = current_section  # Set the last table name as the latest header
        elif para.text.strip():
            extracted_data["content"][current_section] += (" " + para.text.strip()) if current_section else para.text.strip()
        
        # Check if the next element is a table
        if i + 1 < len(paragraphs) and any(table._element.getprevious() is para._element for table in tables):
            table_data = [[cell.text.strip() for cell in row.cells] for row in tables[table_index].rows]
            extracted_data["tables"].append({"table_name": last_table_name, "table_data": table_data})
            table_index += 1  # Move to the next table
    
    extracted_data["content"] = dict(extracted_data["content"])
    return extracted_data

def save_to_json(data, output_file):
    with open(output_file, "w", encoding="utf-8") as json_file:
        json.dump(data, json_file, indent=4, ensure_ascii=False)

if __name__ == "__main__":
    input_file = "data_001.docx"
    output_file = "output.json"
    extracted_data = extract_data_from_docx(input_file)
    save_to_json(extracted_data, output_file)
    print(f"Extraction complete. Data saved to {output_file}")